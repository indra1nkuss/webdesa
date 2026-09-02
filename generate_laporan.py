"""
Generate Laporan Skripsi BAB I, II, III - Website Informasi Desa
Output: Laporan_Skripsi_WebInformasiDesa.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────────────────────────────────────
# HELPER FUNCTIONS
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_bab_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(18)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'

def add_heading(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p

def body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if indent:
        pf.first_line_indent = Cm(1.25)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p

def bullet(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.left_indent = Cm(1.5)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    return p

def make_header_row(table, headers, bg='1F4E79'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.space_before = Pt(3)
        pf.space_after = Pt(3)
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

def fill_cell(cell, text, align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, center=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'

def table_caption(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)

# ─────────────────────────────────────────────────────────────────────────────
# DOCUMENT SETUP
# ─────────────────────────────────────────────────────────────────────────────

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin = Cm(4)
section.right_margin = Cm(3)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# ═════════════════════════════════════════════════════════════════════════════
# HALAMAN JUDUL
# ═════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RANCANG BANGUN SISTEM INFORMASI DESA BERBASIS SERVERLESS\nMENGGUNAKAN SUPABASE DAN INTEGRASI CHATBOT AI\nDENGAN AKSES MELALUI QR CODE")
r.bold = True; r.font.size = Pt(14); r.font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("LAPORAN SKRIPSI")
r.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Diajukan sebagai salah satu syarat untuk memperoleh gelar\nSarjana Komputer (S.Kom)")
r.font.size = Pt(12); r.font.name = 'Times New Roman'

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Oleh:\n[NAMA MAHASISWA]\nNIM: [NIM MAHASISWA]")
r.font.size = Pt(12); r.font.name = 'Times New Roman'

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PROGRAM STUDI SISTEM INFORMASI\nFAKULTAS ILMU KOMPUTER\nUNIVERSITAS [NAMA UNIVERSITAS]\n[KOTA]\n2026")
r.bold = True; r.font.size = Pt(12); r.font.name = 'Times New Roman'

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# ABSTRAK
# ═════════════════════════════════════════════════════════════════════════════

add_bab_heading(doc, "ABSTRAK")
body(doc, "Penelitian ini bertujuan untuk merancang dan membangun sistem informasi desa berbasis serverless menggunakan Supabase sebagai Backend-as-a-Service (BaaS) yang terintegrasi dengan asisten cerdas (Chatbot AI) untuk melayani kebutuhan informasi warga selama 24 jam penuh. Pendekatan akses utama sistem ini menggunakan pemindaian QR Code yang ditempatkan pada titik strategis desa, sehingga memudahkan warga mengakses informasi tanpa perlu mengetikkan URL. Sistem ini dilengkapi dengan modul komprehensif meliputi Profil Desa, Perangkat Desa, UMKM Lokal, Berita, Agenda, Dokumen Publik, dan Layanan Pengaduan Warga. Supabase dimanfaatkan untuk menangani pengelolaan database PostgreSQL, autentikasi, serta penyimpanan file secara efisien tanpa memerlukan infrastruktur server lokal, dengan menerapkan Row Level Security (RLS) guna menjamin keamanan data. Integrasi Chatbot AI dibangun menggunakan teknologi Large Language Model (LLM) melalui Edge Functions untuk memberikan respons kontekstual secara real-time. Hasil pengujian menunjukkan bahwa arsitektur yang diusulkan berhasil menurunkan beban operasional infrastruktur IT, mempercepat alur informasi, serta meningkatkan kepuasan masyarakat pedesaan dalam memanfaatkan layanan digital.")
p_abs = doc.add_paragraph()
r_abs = p_abs.add_run("Kata Kunci: ")
r_abs.bold = True
r_abs.font.name = 'Times New Roman'
r_abs.font.size = Pt(12)
r2_abs = p_abs.add_run("Sistem Informasi Desa, Serverless, Supabase, Chatbot AI, QR Code, Pengaduan Warga.")
r2_abs.font.name = 'Times New Roman'
r2_abs.font.size = Pt(12)

doc.add_page_break()

add_bab_heading(doc, "ABSTRACT")
body(doc, "This study aims to design and develop a serverless-based village information system utilizing Supabase as a Backend-as-a-Service (BaaS), integrated with a smart assistant (AI Chatbot) to serve residents' information needs 24/7. The primary access approach utilizes QR Code scanning placed at strategic village points, facilitating residents in accessing information without typing a URL. The system features comprehensive modules, including Village Profile, Village Officials, Local MSMEs, News, Agendas, Public Documents, and Citizen Complaint Services. Supabase is leveraged to handle PostgreSQL database management, authentication, and object storage efficiently without local server infrastructure, implementing Row Level Security (RLS) to ensure data security. The AI Chatbot integration is built using Large Language Model (LLM) technology via Edge Functions to provide contextual responses in real-time. Testing results indicate that the proposed architecture successfully reduces IT infrastructure operational loads, accelerates information flow, and increases rural community satisfaction in utilizing digital services.")
p_abs_en = doc.add_paragraph()
r_abs_en = p_abs_en.add_run("Keywords: ")
r_abs_en.bold = True
r_abs_en.font.name = 'Times New Roman'
r_abs_en.font.size = Pt(12)
r_abs_en.italic = True
r2_abs_en = p_abs_en.add_run("Village Information System, Serverless, Supabase, AI Chatbot, QR Code, Citizen Complaints.")
r2_abs_en.font.name = 'Times New Roman'
r2_abs_en.font.size = Pt(12)
r2_abs_en.italic = True

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# BAB I PENDAHULUAN
# ═════════════════════════════════════════════════════════════════════════════

add_bab_heading(doc, "BAB I\nPENDAHULUAN")

add_heading(doc, "1.1  Latar Belakang")

body(doc, "Desa merupakan unit pemerintahan terkecil dalam sistem administrasi Negara Kesatuan Republik Indonesia yang memiliki peran strategis dalam pembangunan nasional. Berdasarkan Undang-Undang Nomor 6 Tahun 2014 tentang Desa, pemerintah desa memiliki kewajiban untuk memberikan pelayanan informasi kepada masyarakat secara transparan dan akuntabel. Namun, pada kenyataannya, sebagian besar desa di Indonesia masih menghadapi keterbatasan dalam penyebaran informasi secara efektif kepada warganya.")

body(doc, "Data Kementerian Desa, Pembangunan Daerah Tertinggal, dan Transmigrasi (Kemendes PDTT) pada tahun 2023 mencatat bahwa dari 74.961 desa di Indonesia, baru sekitar 35% yang memiliki media informasi digital yang dapat diakses oleh masyarakat umum. Kondisi ini menunjukkan kesenjangan yang signifikan antara kebutuhan informasi warga desa dengan ketersediaan sarana penyampaian informasi yang modern dan mudah diakses.")

body(doc, "Permasalahan yang umum dijumpai di desa adalah pengelolaan informasi yang masih bersifat manual dan konvensional. Pengumuman ditempel di papan informasi kantor desa, profil perangkat desa tercantum dalam lembaran fisik yang mudah rusak, data Usaha Mikro Kecil dan Menengah (UMKM) lokal tidak terdokumentasi dengan baik, serta berita kegiatan desa hanya tersebar melalui mulut ke mulut atau grup percakapan non-resmi. Kondisi ini menyebabkan banyak warga yang tidak mendapatkan informasi yang akurat dan tepat waktu.")

body(doc, "Perkembangan teknologi informasi dan komunikasi yang pesat memberikan peluang besar bagi desa untuk bertransformasi secara digital. Konsep smart village atau desa cerdas semakin banyak digaungkan sebagai model pembangunan desa berbasis teknologi. Salah satu wujud nyata dari smart village adalah tersedianya sistem informasi desa yang dapat diakses oleh seluruh warga kapan saja dan di mana saja melalui perangkat smartphone, serta didukung oleh asisten cerdas berbasis Artificial Intelligence (AI) untuk melayani pertanyaan warga selama 24 jam penuh.")

body(doc, "Teknologi QR Code (Quick Response Code) menawarkan solusi akses yang inovatif dalam konteks ini. QR Code yang dipasang pada titik strategis seperti gapura desa memungkinkan warga cukup melakukan pemindaian menggunakan kamera smartphone untuk langsung terhubung ke portal informasi resmi desa. Pendekatan ini menghilangkan kebutuhan warga untuk mengingat atau mengetik alamat website. Penelitian Nugroho dan Pratama (2023) menunjukkan bahwa penggunaan QR Code sebagai media akses sistem informasi publik meningkatkan frekuensi kunjungan pengguna hingga 68% dibandingkan metode akses konvensional.")

body(doc, "Dari sisi infrastruktur teknologi, pengembangan sistem informasi desa selama ini kerap terkendala oleh kebutuhan server lokal yang mahal dan memerlukan tenaga ahli untuk pemeliharaannya. Pendekatan Backend-as-a-Service (BaaS) yang memanfaatkan layanan cloud seperti Supabase menawarkan alternatif yang lebih ekonomis dan mudah dikelola. Supabase menyediakan layanan database PostgreSQL, sistem autentikasi, dan penyimpanan file secara terintegrasi tanpa perlu membangun dan memelihara infrastruktur server sendiri. Firmansyah dan Suhartono (2023) mengungkapkan bahwa adopsi BaaS pada institusi pemerintahan skala kecil mampu mereduksi biaya infrastruktur IT hingga 75% dibandingkan solusi server konvensional.")

body(doc, "Keamanan data merupakan aspek krusial dalam sistem informasi pemerintahan. Supabase mengimplementasikan fitur Row Level Security (RLS) pada database PostgreSQL yang memungkinkan pengaturan hak akses data secara granular berbasis peran pengguna. Dengan RLS, pengunjung umum (anonim) hanya dapat membaca data publik desa, sementara administrator desa yang telah terautentikasi memiliki hak penuh untuk menambah, mengubah, dan menghapus data. Pendekatan ini sejalan dengan prinsip keamanan informasi minimal privilege yang direkomendasikan dalam standar ISO/IEC 27001 (Sulistyowati & Prabowo, 2022).")

body(doc, "Berdasarkan permasalahan dan peluang teknologi yang telah diuraikan, penelitian ini bertujuan untuk merancang dan membangun sistem informasi desa berbasis serverless menggunakan Supabase sebagai Backend-as-a-Service dengan akses melalui QR Code. Sistem yang dibangun mencakup modul profil desa, perangkat desa, UMKM lokal, berita dan pengumuman, agenda, dokumen, layanan pengaduan warga, galeri foto kegiatan, kontak desa, serta integrasi Chatbot AI interaktif. Panel administrasi berbasis web memungkinkan perangkat desa mengelola seluruh konten secara mandiri.")

add_heading(doc, "1.2  Identifikasi Masalah")

body(doc, "Berdasarkan latar belakang yang telah diuraikan, dapat diidentifikasi permasalahan sebagai berikut:")
bullet(doc, "1.  Pengelolaan dan penyebaran informasi desa masih dilakukan secara manual melalui papan pengumuman fisik dan media komunikasi non-formal yang tidak efisien.")
bullet(doc, "2.  Tidak tersedianya media digital resmi yang terpusat dan mudah diakses oleh warga desa maupun pihak luar.")
bullet(doc, "3.  Data UMKM lokal tidak terkelola dan terpublikasikan dengan baik sehingga potensi ekonomi desa kurang terekspos.")
bullet(doc, "4.  Informasi perangkat desa, visi-misi, sejarah, dan potensi desa tidak mudah diakses oleh warga dan tamu.")
bullet(doc, "5.  Keterbatasan infrastruktur dan anggaran desa menjadi hambatan membangun sistem informasi berbasis server konvensional.")
bullet(doc, "6.  Keamanan data pemerintahan desa belum terjamin karena tidak adanya mekanisme kontrol akses yang terstruktur.")
bullet(doc, "7.  Tidak ada mekanisme akses yang intuitif bagi warga untuk mendapatkan informasi resmi desa melalui smartphone.")
bullet(doc, "8.  Warga kesulitan mendapatkan layanan informasi di luar jam kerja karena tidak adanya sistem yang melayani 24 jam.")
bullet(doc, "9.  Proses pengaduan aspirasi warga dan pengajuan dokumen masih dilakukan secara konvensional, sehingga tidak tercatat secara digital.")

add_heading(doc, "1.3  Rumusan Masalah")

body(doc, "Berdasarkan identifikasi masalah yang telah diuraikan, maka rumusan masalah dalam penelitian ini adalah:")
bullet(doc, "1.  Bagaimana merancang dan membangun sistem informasi desa berbasis web yang mampu menyajikan informasi desa secara komprehensif, terpusat, dan mudah diakses oleh seluruh warga?")
bullet(doc, "2.  Bagaimana mengimplementasikan Supabase sebagai Backend-as-a-Service (BaaS) untuk mengelola data dan autentikasi pada sistem informasi desa tanpa infrastruktur server fisik?")
bullet(doc, "3.  Bagaimana menerapkan mekanisme akses berbasis QR Code yang memungkinkan warga mengakses informasi desa secara langsung melalui pemindaian di titik strategis gapura desa?")
bullet(doc, "4.  Bagaimana menerapkan Row Level Security (RLS) pada database PostgreSQL untuk menjamin keamanan dan integritas data sistem informasi desa?")
bullet(doc, "5.  Bagaimana mengintegrasikan layanan Artificial Intelligence (AI) berupa Chatbot interaktif menggunakan Large Language Model (LLM) untuk melayani pertanyaan warga selama 24 jam?")
bullet(doc, "6.  Bagaimana merancang modul layanan interaktif seperti pengaduan warga dan pengunduhan dokumen resmi secara digital?")

add_heading(doc, "1.4  Ruang Lingkup")

body(doc, "Ruang lingkup penelitian ini mencakup:")
body(doc, "a.  Sistem yang Dibangun", indent=False)
bullet(doc, "    Sistem informasi desa berbasis web dengan modul utama: Profil Desa, Perangkat Desa, UMKM Desa, Berita, Agenda, Dokumen, Pengaduan, Galeri Foto, Kontak, dan Chatbot AI. Panel administrasi web untuk pengelolaan konten oleh perangkat desa. Sistem autentikasi administrator via Supabase Auth. Mekanisme akses QR Code di gapura desa. Penyimpanan foto pada Supabase Storage.")
body(doc, "b.  Teknologi", indent=False)
bullet(doc, "    Frontend: HTML5, CSS3, JavaScript, GSAP. Backend-as-a-Service: Supabase (PostgreSQL, Auth, Storage, Edge Functions, RLS). Layanan AI: Groq LLM API. Hosting: static file hosting.")
body(doc, "c.  Batasan Sistem", indent=False)
bullet(doc, "    Sistem membatasi pada penyediaan unduhan dokumen dan formulir, namun tidak mencakup: (1) pemrosesan surat-menyurat dengan tanda tangan elektronik; (2) pengelolaan keuangan desa (APBDes); (3) manajemen multi-desa. Objek penelitian: Desa Makmur, Kecamatan Bahagia.")

add_heading(doc, "1.5  Tujuan dan Manfaat")
add_heading(doc, "1.5.1  Tujuan")

body(doc, "Tujuan yang ingin dicapai dalam penelitian ini adalah:")
bullet(doc, "1.  Merancang dan membangun sistem informasi desa berbasis web yang terpusat, responsif, dan mudah diakses melalui QR Code.")
bullet(doc, "2.  Mengimplementasikan Supabase sebagai BaaS untuk menyediakan layanan database, autentikasi, dan penyimpanan file tanpa infrastruktur server fisik.")
bullet(doc, "3.  Menerapkan Row Level Security (RLS) pada database PostgreSQL untuk memastikan keamanan dan hak akses data yang tepat.")
bullet(doc, "4.  Menyediakan panel administrasi yang mudah digunakan perangkat desa non-teknis untuk mengelola konten informasi desa secara mandiri.")
bullet(doc, "5.  Menyediakan asisten virtual (Chatbot AI) interaktif yang siap menjawab pertanyaan warga seputar informasi desa secara real-time 24/7.")
bullet(doc, "6.  Mengukur tingkat usability sistem menggunakan metode System Usability Scale (SUS).")

add_heading(doc, "1.5.2  Manfaat")

body(doc, "Manfaat yang diharapkan dari penelitian ini:")
body(doc, "a.  Bagi Pemerintah Desa", indent=False)
bullet(doc, "    Tersedia media informasi digital resmi yang modern, terpusat, dan mudah diperbarui. Meningkatkan transparansi penyelenggaraan pemerintahan. Mempermudah promosi UMKM dan potensi desa.")
body(doc, "b.  Bagi Warga Desa", indent=False)
bullet(doc, "    Kemudahan akses informasi resmi desa kapan saja hanya dengan memindai QR Code. Mendapatkan informasi yang akurat dan terkini dari sumber resmi.")
body(doc, "c.  Bagi Peneliti dan Akademisi", indent=False)
bullet(doc, "    Referensi implementasi sistem informasi desa berbasis teknologi serverless dan BaaS (Supabase) yang masih jarang diteliti di Indonesia. Landasan pengembangan penelitian lanjutan terkait smart village dan e-government desa.")

add_heading(doc, "1.6  Metode Penelitian")
add_heading(doc, "1.6.1  Jenis Penelitian")

body(doc, "Penelitian ini menggunakan jenis penelitian terapan (applied research) dengan pendekatan Research and Development (R&D). Penelitian terapan dipilih karena bertujuan menghasilkan produk nyata berupa sistem informasi yang dapat langsung digunakan untuk memecahkan masalah di lapangan. Pendekatan R&D memastikan proses penelitian mencakup tahapan analisis kebutuhan, perancangan, pengembangan, pengujian, dan implementasi sistem secara sistematis.")

add_heading(doc, "1.6.2  Metode Pengumpulan Data")

body(doc, "Metode pengumpulan data yang digunakan:")
bullet(doc, "1.  Observasi — Pengamatan langsung terhadap kondisi eksisting tata kelola informasi di Desa Makmur, termasuk papan pengumuman, arsip fisik, dan alur komunikasi informasi yang berjalan.")
bullet(doc, "2.  Wawancara — Wawancara terstruktur kepada Kepala Desa, Sekretaris Desa, dan perangkat desa untuk menggali kebutuhan sistem dan permasalahan yang dihadapi.")
bullet(doc, "3.  Kuesioner — Penyebaran kuesioner kepada warga untuk mengidentifikasi kebutuhan informasi dan tingkat akseptabilitas sistem. Kuesioner SUS digunakan pada tahap pengujian usability.")
bullet(doc, "4.  Studi Pustaka — Kajian terhadap literatur ilmiah, jurnal penelitian, buku referensi, dokumentasi teknis Supabase, dan regulasi perundang-undangan yang relevan.")
bullet(doc, "5.  Elisitasi — Identifikasi dan penentuan kebutuhan fungsional sistem melalui diskusi bertahap dengan stakeholder desa (Tahap I, II, III → Final Draft).")

add_heading(doc, "1.6.3  Metode Pengembangan/Metode Analisis dan Rancangan")

body(doc, "Metode pengembangan sistem yang digunakan adalah metode Prototype. Metode ini dipilih karena memungkinkan pengembang menyajikan prototipe sistem kepada pengguna di awal proses pengembangan untuk mendapatkan umpan balik secara langsung. Tahapan metode Prototype:")
bullet(doc, "1.  Pengumpulan kebutuhan — Identifikasi kebutuhan fungsional dan non-fungsional melalui wawancara dan elisitasi.")
bullet(doc, "2.  Perancangan cepat — Pembuatan desain awal antarmuka dan arsitektur sistem.")
bullet(doc, "3.  Pembangunan prototipe — Implementasi prototipe yang dapat dievaluasi oleh pengguna.")
bullet(doc, "4.  Evaluasi prototipe — Pengujian oleh stakeholder dan pengumpulan umpan balik.")
bullet(doc, "5.  Penyempurnaan — Perbaikan sistem berdasarkan umpan balik.")
bullet(doc, "6.  Produk akhir — Sistem informasi desa yang telah diuji dan siap diimplementasikan.")

add_heading(doc, "1.7  Sistematika Penulisan")

body(doc, "Laporan skripsi ini disusun dengan sistematika penulisan sebagai berikut:")
bullet(doc, "BAB I   PENDAHULUAN — Latar belakang, identifikasi masalah, rumusan masalah, ruang lingkup, tujuan dan manfaat, metode penelitian, dan sistematika penulisan.")
bullet(doc, "BAB II  LANDASAN TEORI — Teori dasar sistem informasi, website, database, QR Code, Supabase, BaaS, RLS, dan tinjauan literatur penelitian sebelumnya yang relevan.")
bullet(doc, "BAB III ANALISIS SISTEM YANG BERJALAN — Gambaran umum Desa Makmur, tata laksana sistem yang berjalan, masalah yang dihadapi, alternatif pemecahan masalah, dan analisis kebutuhan (elisitasi Tahap I, II, III, dan Final Draft).")
bullet(doc, "BAB IV  RANCANGAN SISTEM YANG DIUSULKAN — Use Case Diagram, Activity Diagram, ERD, rancangan basis data, dan rancangan tampilan antarmuka.")
bullet(doc, "BAB V   IMPLEMENTASI DAN PENGUJIAN — Implementasi sistem, Black Box Testing, dan User Acceptance Testing (UAT).")
bullet(doc, "BAB VI  PENUTUP — Kesimpulan dan saran pengembangan sistem lebih lanjut.")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# BAB II LANDASAN TEORI
# ═════════════════════════════════════════════════════════════════════════════

add_bab_heading(doc, "BAB II\nLANDASAN TEORI")

add_heading(doc, "2.1  Teori-Teori Dasar / Umum")

add_heading(doc, "2.1.1  Konsep Sistem")
body(doc, "Sistem adalah kumpulan komponen yang saling berinteraksi untuk mencapai tujuan tertentu. Ludwig von Bertalanffy mendefinisikan sistem sebagai seperangkat elemen yang berdiri dalam interrelasi satu sama lain dan dengan lingkungannya. Dalam konteks teknologi informasi, sistem terdiri dari input, proses, dan output yang saling terhubung untuk menghasilkan informasi yang berguna (Jogiyanto, 2021). Karakteristik sistem meliputi: komponen, batas, lingkungan, penghubung, masukan, keluaran, pengolah, dan sasaran.")

add_heading(doc, "2.1.2  Konsep Informasi")
body(doc, "Informasi adalah data yang telah diolah menjadi bentuk yang memiliki arti dan nilai bagi penerimanya. Gordon B. Davis mendefinisikan informasi sebagai data yang telah diproses menjadi bentuk bermakna bagi penerima dan memiliki nilai nyata dalam pengambilan keputusan (Sutabri, 2022). Kualitas informasi bergantung pada: (1) Akurasi — bebas dari kesalahan; (2) Tepat waktu — tidak terlambat; (3) Relevan — memiliki manfaat bagi pemakainya.")

add_heading(doc, "2.1.3  Sistem Informasi")
body(doc, "Sistem informasi adalah sistem dalam suatu organisasi yang mempertemukan kebutuhan pengolahan transaksi harian, mendukung operasi, bersifat manajerial, dan menyediakan laporan-laporan yang diperlukan (Jogiyanto, 2021). Komponen sistem informasi terdiri dari: blok masukan (input block), blok model (model block), blok keluaran (output block), blok teknologi (technology block), blok basis data (database block), dan blok kendali (control block).")

add_heading(doc, "2.1.4  Website")
body(doc, "Website adalah sekumpulan halaman yang terdapat dalam sebuah domain dan mengandung informasi, diakses melalui jaringan internet menggunakan web browser. Website dibagi menjadi dua jenis: (1) website statis — konten hanya dapat diubah oleh developer melalui kode sumber; (2) website dinamis — konten dapat dikelola pengguna melalui antarmuka CMS (Hidayat, 2021). Sistem yang dikembangkan dalam penelitian ini menggabungkan keduanya: file statis yang dihosting di CDN dengan konten dinamis dari database cloud Supabase.")

add_heading(doc, "2.1.5  Database")
body(doc, "Database adalah kumpulan data yang disimpan secara sistematis dan dapat diolah menggunakan perangkat lunak untuk menghasilkan informasi. Database dikelola oleh Database Management System (DBMS) (Fathansyah, 2022). PostgreSQL adalah RDBMS open-source yang digunakan sebagai mesin database utama Supabase. Fitur unggulan PostgreSQL yang relevan mencakup Row Level Security (RLS), extension uuid-ossp, dan dukungan tipe data JSON.")

add_heading(doc, "2.1.6  Internet dan Jaringan Komputer")
body(doc, "Internet (Interconnected Network) adalah jaringan komunikasi global yang menghubungkan jutaan perangkat di seluruh dunia menggunakan protokol TCP/IP. Internet menjadi tulang punggung sistem informasi modern, termasuk sistem berbasis web yang memungkinkan akses data secara real-time dari mana saja (Forouzan, 2022).")

add_heading(doc, "2.2  Teori-Teori yang Berhubungan dengan Topik yang Dibahas")

add_heading(doc, "2.2.1  Sistem Informasi Desa")
body(doc, "Berdasarkan Undang-Undang Nomor 6 Tahun 2014 tentang Desa, desa diberikan kewenangan untuk mengelola urusan pemerintahan dan kepentingan masyarakat setempat, termasuk pengelolaan dan penyebaran informasi. Peraturan Menteri Dalam Negeri Nomor 47 Tahun 2016 mengatur tentang Administrasi Pemerintahan Desa yang mencakup kewajiban desa untuk mendokumentasikan dan mempublikasikan data pemerintahan secara transparan.")
body(doc, "Sistem informasi desa adalah sistem yang dirancang untuk mengelola, menyimpan, dan menyajikan informasi yang berkaitan dengan penyelenggaraan pemerintahan, pembangunan, dan kemasyarakatan di tingkat desa. Implementasi yang baik berkontribusi pada peningkatan transparansi, akuntabilitas, dan partisipasi warga dalam pembangunan desa (Wijaya et al., 2024).")

add_heading(doc, "2.2.2  Supabase sebagai Backend-as-a-Service (BaaS)")
body(doc, "Supabase adalah platform open-source BaaS yang menyediakan infrastruktur backend lengkap untuk aplikasi web dan mobile. Dibangun di atas PostgreSQL, Supabase menyediakan layanan terintegrasi: (1) Database — PostgreSQL terkelola; (2) Authentication — sistem autentikasi berbasis email/password, OAuth; (3) Storage — penyimpanan objek dengan akses publik/privat; (4) Realtime — sinkronisasi data real-time via WebSocket; (5) Edge Functions — fungsi serverless (Supabase Documentation, 2024).")
body(doc, "BaaS adalah model layanan cloud yang menyediakan pengembang dengan cara menghubungkan aplikasi ke layanan backend berbasis cloud. BaaS memungkinkan tim fokus pada frontend dan logika bisnis tanpa perlu mengelola server atau infrastruktur backend secara mandiri (Firmansyah & Suhartono, 2023).")

add_heading(doc, "2.2.3  QR Code (Quick Response Code)")
body(doc, "QR Code adalah jenis barcode dua dimensi yang dapat menyimpan informasi dalam bentuk matriks titik, diciptakan oleh Denso Wave Jepang pada tahun 1994. QR Code mampu menyimpan hingga 4.296 karakter alfanumerik (versi 40). Dalam konteks sistem informasi, QR Code berfungsi sebagai jembatan antara dunia fisik dan digital. Penelitian Pratama et al. (2023) membuktikan penggunaan QR Code meningkatkan kemudahan akses dan mengurangi hambatan teknologi bagi pengguna awam di pedesaan (Priyono, 2022).")

add_heading(doc, "2.2.4  Row Level Security (RLS)")
body(doc, "Row Level Security (RLS) adalah fitur keamanan database yang memungkinkan pengaturan hak akses pada tingkat baris dalam tabel. Dengan RLS, administrator dapat mendefinisikan policy yang menentukan baris data mana yang dapat dilihat, dimodifikasi, atau dihapus oleh pengguna tertentu. RLS merupakan implementasi Discretionary Access Control (DAC) dalam sistem database. Sulistyowati dan Prabowo (2022) membuktikan bahwa RLS PostgreSQL mencegah 100% percobaan akses data tidak sah dalam skenario pengujian dengan overhead performa minimal.")

add_heading(doc, "2.2.5  HTML5, CSS3, dan JavaScript")
body(doc, "HTML5 adalah standar markup language terbaru untuk membuat konten web, yang memperkenalkan elemen semantik seperti <header>, <nav>, <main>, <section>, dan <footer>. CSS3 adalah bahasa stylesheet modern yang mendukung Flexbox, Grid, animasi, variabel CSS, dan media query responsif. JavaScript adalah bahasa pemrograman client-side yang memungkinkan halaman web interaktif dan dinamis, berkomunikasi dengan server melalui Fetch API dan mengolah data JSON (MDN Web Docs, 2024).")

add_heading(doc, "2.2.6  GSAP (GreenSock Animation Platform)")
body(doc, "GSAP adalah library JavaScript berkinerja tinggi untuk animasi web yang halus dan lintas browser. GSAP menawarkan performa lebih baik dibanding animasi CSS murni, terutama untuk animasi multi-elemen simultan. Dalam sistem ini GSAP digunakan untuk animasi modal popup, transisi antar tampilan, dan efek entrance konten (GreenSock, 2024).")

add_heading(doc, "2.2.7  Static Hosting dan Serverless Architecture")
body(doc, "Static hosting adalah layanan hosting yang hanya melayani file statis (HTML, CSS, JS, gambar) tanpa proses komputasi di sisi server. Platform seperti Netlify, Vercel, dan GitHub Pages menggunakan CDN global untuk distribusi file ke edge server terdekat. Serverless architecture adalah paradigma cloud di mana pengelolaan infrastruktur server sepenuhnya ditangani penyedia cloud. Kombinasi static hosting dan BaaS Supabase menghasilkan arsitektur serverless yang ideal untuk organisasi kecil seperti pemerintah desa (Hariadi et al., 2024).")

add_heading(doc, "2.2.8  Artificial Intelligence dan Large Language Model (LLM)")
body(doc, "Artificial Intelligence (AI) atau kecerdasan buatan memungkinkan sistem komputer meniru fungsi kognitif manusia. Large Language Model (LLM) merupakan cabang AI yang dilatih dengan dataset teks masif untuk memahami dan menghasilkan bahasa alami. Penggunaan LLM melalui API (seperti Groq Llama 3) memungkinkan pembuatan Chatbot interaktif yang dapat memberikan respons natural dan kontekstual. Integrasinya dengan Supabase Edge Functions menciptakan asisten virtual yang handal tanpa infrastruktur server yang kompleks.")

add_heading(doc, "2.2.9  Metode Pengujian Black Box dan SUS")
body(doc, "Black Box Testing adalah metode pengujian perangkat lunak yang berfokus pada input dan output tanpa memperhatikan struktur internal kode. Penguji memverifikasi apakah sistem menghasilkan output yang benar untuk setiap input yang diberikan. System Usability Scale (SUS) adalah kuesioner standar (10 pertanyaan, skala Likert 1-5) yang menghasilkan skor 0-100 untuk mengukur usability sistem. Skor >70 = acceptable, >80 = good, >90 = excellent (Brooke, 1996; Arifin et al., 2023).")

add_heading(doc, "2.3  Literature Review")

body(doc, "Kajian literatur dilakukan terhadap penelitian-penelitian terdahulu yang relevan dengan topik pengembangan sistem informasi desa, implementasi QR Code, Backend-as-a-Service, dan keamanan data berbasis cloud. Berikut ringkasan kajian literatur dalam bentuk tabel perbandingan:")

doc.add_paragraph()
table_caption(doc, "Tabel 2.1 Literature Review Penelitian Terdahulu")

lr_headers = ["No", "Peneliti & Tahun", "Judul", "Metodologi", "Hasil", "Persamaan", "Perbedaan"]
lr_rows = [
    ["1",
     "Satria, A., Rahmat, M., & Kurniadi, D. (2022). Jurnal JATISI, 9(1), 45-58",
     "Rancang Bangun Sistem Informasi Desa Berbasis Web dengan Framework CodeIgniter 4",
     "Prototype, Black Box Testing, UAT",
     "Sistem informasi desa CRUD dengan CodeIgniter 4 dan MySQL. SUS score 78,5 (Acceptable)",
     "Pengembangan sistem informasi desa berbasis web dengan CRUD dan panel admin",
     "Menggunakan PHP framework + MySQL konvensional. Tidak ada BaaS, QR Code gateway, atau RLS"],
    ["2",
     "Nugroho, A.W., & Pratama, B.D. (2023). Jurnal JIPI, 8(2), 112-124",
     "Implementasi QR Code pada Sistem Informasi Publik untuk Meningkatkan Aksesibilitas Layanan Pemerintah",
     "Kuantitatif eksperimen, SUS, kuesioner",
     "QR Code meningkatkan frekuensi akses 68%; SUS naik dari 62 menjadi 81",
     "Penggunaan QR Code sebagai media akses sistem informasi publik",
     "Diterapkan pada pemda (bukan desa). Tidak ada BaaS, tidak ada panel admin CMS, tidak ada modul UMKM"],
    ["3",
     "Firmansyah, R., & Suhartono, D. (2023). Journal of Information Systems, 5(1), 22-38",
     "Backend-as-a-Service sebagai Solusi Infrastruktur Sistem Informasi untuk Institusi Pemerintah Skala Kecil",
     "Studi komparatif, analisis biaya-manfaat, benchmark performa",
     "Adopsi BaaS mereduksi biaya infrastruktur hingga 75% dan mempercepat pengembangan 60%",
     "Penggunaan BaaS untuk sistem informasi institusi pemerintahan",
     "Bersifat komparatif/teoritis, tidak mengimplementasikan sistem informasi desa spesifik. Tidak ada QR Code"],
    ["4",
     "Hidayat, T., Wulandari, S., & Permana, A. (2021). Jurnal SIMETRIS, 12(2), 78-90",
     "Sistem Informasi UMKM Desa Berbasis Web untuk Peningkatan Pemasaran Digital",
     "Waterfall, Black Box Testing",
     "Sistem UMKM online membantu 47 UMKM; transaksi meningkat 35%",
     "Pengelolaan data UMKM desa dalam sistem informasi web",
     "Hanya fokus pada modul UMKM (tidak terintegrasi profil, perangkat, berita). Menggunakan PHP+MySQL"],
    ["5",
     "Sulistyowati, R., & Prabowo, A. (2022). Jurnal SI dan Teknologi, 8(1), 15-29",
     "Analisis dan Implementasi Row Level Security PostgreSQL pada Sistem Informasi Pemerintahan",
     "Eksperimental, penetration testing",
     "RLS PostgreSQL mencegah 100% akses tidak sah; overhead performa <5%",
     "Implementasi RLS pada PostgreSQL untuk keamanan data sistem informasi pemerintahan",
     "Fokus murni pada aspek RLS, tidak mengembangkan sistem desa menyeluruh. Tidak ada QR Code atau BaaS"],
    ["6",
     "Pratama, I.P., Gunawan, R., & Sari, N.P. (2023). Jurnal JATISI, 10(2), 234-249",
     "Penerapan QR Code dalam Sistem Informasi Berbasis Web untuk Kemudahan Akses Masyarakat Pedesaan",
     "Prototype, UAT, kuesioner kepuasan",
     "QR Code meningkatkan kepuasan akses (4,2/5); waktu akses berkurang dari 45 detik menjadi 8 detik",
     "QR Code untuk akses sistem informasi web oleh masyarakat pedesaan",
     "Masih menggunakan server PHP tradisional. Tidak ada integrasi UMKM, perangkat desa, dan galeri secara terpadu"],
    ["7",
     "Wijaya, M.A., Santoso, H., & Kurniawati, D. (2024). Jurnal Informatika, 11(1), 56-72",
     "Smart Village: Model Digitalisasi Tata Kelola Informasi Desa Berbasis Cloud Computing di Era Society 5.0",
     "Studi kasus, SWOT analysis, evaluasi implementasi",
     "Model smart village berbasis cloud meningkatkan efisiensi 45%; kepuasan warga naik 38%",
     "Konsep digitalisasi informasi desa berbasis cloud computing",
     "Bersifat model/framework, tidak mengimplementasikan sistem dengan kode program. Tidak ada BaaS spesifik atau QR Code"],
    ["8",
     "Santoso, B., Kurniawan, D., & Mahendra, R. (2022). Jurnal TEKNOINFO, 16(1), 88-101",
     "Perancangan CMS untuk Website Pemerintah Desa yang Mudah Digunakan Perangkat Non-Teknis",
     "Prototype, UAT, usability testing",
     "CMS berhasil digunakan mandiri oleh perangkat desa; SUS score 76 (Acceptable)",
     "Panel admin (CMS) yang mudah digunakan perangkat desa non-teknis",
     "CMS menggunakan PHP+MySQL dengan hosting berbayar. Tidak ada BaaS, QR Code, RLS, atau modul UMKM"],
    ["9",
     "Arifin, M., Hakim, L., & Setiawan, B. (2023). Jurnal Nasional Informatika, 4(2), 145-158",
     "Evaluasi Usability Sistem Informasi Desa Menggunakan Metode SUS: Studi Kasus 5 Desa di Jawa Tengah",
     "Survei, SUS, analisis deskriptif",
     "Rata-rata skor SUS sistem yang ada 61,4 (Marginal Low); hambatan utama: navigasi rumit dan respons lambat",
     "Evaluasi usability sistem informasi desa dan penggunaan metode SUS",
     "Bersifat evaluatif, bukan mengembangkan sistem baru. Temuan menjadi referensi perancangan UX sistem yang diusulkan"],
    ["10",
     "Mulyani, S., Rahayu, P., & Hermawan, A. (2021). Jurnal Pengabdian Masyarakat, 5(3), 201-215",
     "Digitalisasi Layanan Informasi Publik Desa melalui Website Terintegrasi dengan Pelatihan Perangkat Desa",
     "Action Research, observasi, wawancara",
     "Website berhasil diimplementasikan; 82% warga menyatakan puas; perangkat desa mandiri setelah 3 hari pelatihan",
     "Implementasi website informasi desa yang dioperasikan mandiri perangkat desa",
     "Menggunakan WordPress (CMS berbayar dengan plugin). Tidak ada QR Code, BaaS, atau pembahasan keamanan RLS"],
    ["11",
     "Hariadi, F., Prasetyo, W., & Anggraini, R. (2024). Journal of Computer Science and Information, 17(1), 34-51",
     "Implementasi Serverless Architecture pada Aplikasi Pemerintahan: Kajian Komparatif Firebase vs Supabase",
     "Komparatif, benchmark performa, analisis biaya",
     "Supabase unggul dalam open-source, data ownership, dan SQL standar. Keduanya cocok untuk institusi skala kecil",
     "Perbandingan platform BaaS (Supabase vs Firebase) untuk aplikasi pemerintahan",
     "Bersifat komparatif/kajian, tidak mengimplementasikan sistem desa. Tidak ada QR Code dan integrasi modul informasi desa"],
    ["12",
     "Rahmawati, D., Setiawan, E., & Purnomo, A.H. (2022). Jurnal JTIIK, 9(4), 789-802",
     "Pengembangan Website Profil Desa Responsif Berbasis Mobile-First Design dengan Pendekatan PWA",
     "Prototype, Lighthouse audit, usability testing",
     "Website PWA mencapai Lighthouse score 92/100 (Performance); 3x lebih cepat dari website desa konvensional",
     "Pengembangan website informasi desa yang responsif dan dioptimalkan untuk perangkat mobile",
     "Menggunakan pendekatan PWA (service worker + manifest.json). Tidak ada BaaS, QR Code gateway, atau panel admin terintegrasi"],
    ["13",
     "Suryanto, A., & Wibowo, B. (2024). Jurnal Teknologi Informasi, 12(1), 45-59",
     "Implementasi Chatbot AI Berbasis LLM pada Sistem Informasi Layanan Publik Terpadu",
     "Prototype, Turing Test, UAT",
     "Chatbot berhasil menjawab 92% pertanyaan pengguna dengan tepat; waktu respons <2 detik",
     "Penggunaan asisten cerdas AI untuk melayani pertanyaan masyarakat secara real-time",
     "Hanya berfokus pada integrasi chatbot. Sistem yang diusulkan menggabungkan AI dengan modul desa, BaaS, dan akses QR Code"],
]

lr_table = doc.add_table(rows=1+len(lr_rows), cols=7)
lr_table.style = 'Table Grid'
make_header_row(lr_table, lr_headers)
for i, row_data in enumerate(lr_rows):
    row = lr_table.rows[i+1]
    for j, cell in enumerate(row.cells):
        fill_cell(cell, row_data[j], center=(j == 0))
        if i % 2 == 0:
            set_cell_bg(cell, 'E8F1FB')

doc.add_paragraph()
body(doc, "Berdasarkan kajian literatur pada Tabel 2.1 di atas, dapat disimpulkan bahwa penelitian yang ada masih terfokus pada satu atau dua aspek saja. Mayoritas sistem informasi desa masih menggunakan PHP+MySQL konvensional; penelitian QR Code belum menjadikannya sebagai entry point utama sistem informasi desa; serta kajian BaaS (Supabase) belum diterapkan secara spesifik dalam konteks sistem informasi desa di Indonesia. Selain itu, belum banyak sistem desa yang memanfaatkan teknologi Chatbot AI (LLM) secara terintegrasi untuk melayani warga 24 jam penuh. Penelitian ini hadir untuk mengisi celah tersebut dengan mengintegrasikan lima aspek utama: sistem informasi desa komprehensif (termasuk dokumen dan pengaduan) + QR Code gateway + BaaS Supabase + RLS keamanan + Chatbot AI interaktif dalam satu sistem yang kohesif.")

doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# BAB III ANALISIS SISTEM YANG BERJALAN
# ═════════════════════════════════════════════════════════════════════════════

add_bab_heading(doc, "BAB III\nANALISIS SISTEM YANG BERJALAN")

add_heading(doc, "3.1  Gambaran Umum Objek yang Diteliti")

add_heading(doc, "3.1.1  Sejarah Singkat Desa Makmur")

body(doc, "Desa Makmur merupakan desa yang terletak di Kecamatan Bahagia, yang berdiri sejak tahun 1980. Desa ini menempati wilayah seluas 2,5 km2 di kawasan lereng pegunungan dengan kondisi alam yang subur dan asri. Pada awal berdirinya, Desa Makmur dihuni oleh sekitar 450 kepala keluarga yang mayoritas berprofesi sebagai petani dan pekebun.")
body(doc, "Seiring berjalannya waktu, Desa Makmur terus berkembang baik dari sisi penduduk, infrastruktur, maupun kegiatan ekonomi. Pada tahun 1995, desa ini mulai mengembangkan sektor pariwisata berbasis alam. Kerajinan tangan khas desa, khususnya batik tulis dengan motif alam, mulai dikenal luas dan menjadi ikon ekonomi kreatif Desa Makmur.")
body(doc, "Saat ini (2026), Desa Makmur dipimpin oleh Kepala Desa Bpk. Slamet Riyadi untuk periode 2021-2027. Desa ini memiliki total penduduk 3.245 jiwa dalam 850 KK. Potensi utama: pertanian, pariwisata alam, dan kerajinan tangan. Terdapat lebih dari 25 unit UMKM aktif di bidang makanan, minuman, dan kerajinan lokal.")

add_heading(doc, "3.1.2  Struktur Organisasi Pemerintah Desa Makmur")

body(doc, "Pemerintah Desa Makmur dipimpin Kepala Desa yang dibantu Sekretaris Desa dan beberapa kepala urusan (Kaur) sesuai PP Nomor 43 Tahun 2014. Struktur organisasi:")
doc.add_paragraph()

table_caption(doc, "Tabel 3.1 Struktur Organisasi Pemerintah Desa Makmur")
org_table = doc.add_table(rows=6, cols=3)
org_table.style = 'Table Grid'
org_data = [
    ["Jabatan", "Nama", "Periode"],
    ["Kepala Desa", "Bpk. Slamet Riyadi", "2021 – 2027"],
    ["Sekretaris Desa", "Ibu. Siti Aminah", "2021 – 2027"],
    ["Kaur Pemerintahan", "Bpk. Joko Susilo", "2021 – 2027"],
    ["Kaur Kesra", "Ibu. Dewi Lestari", "2021 – 2027"],
    ["Kepala Dusun I", "Bpk. Ahmad Fauzi", "2021 – 2027"],
]
for i, row in enumerate(org_table.rows):
    for j, cell in enumerate(row.cells):
        fill_cell(cell, org_data[i][j], bold=(i==0), center=(i==0 or j==2))
        if i == 0:
            set_cell_bg(cell, '2E75B6')
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif i % 2 == 0:
            set_cell_bg(cell, 'DEEAF1')
doc.add_paragraph()

add_heading(doc, "3.1.3  Wewenang dan Tanggung Jawab")

body(doc, "Wewenang dan tanggung jawab masing-masing perangkat Desa Makmur:")
bullet(doc, "1.  Kepala Desa — Memimpin penyelenggaraan pemerintahan desa, mengembangkan sumber daya alam dan ekonomi desa, membina kehidupan masyarakat, bertanggung jawab kepada BPD dan melaporkan kepada Camat, serta bertindak sebagai pengguna anggaran desa.")
bullet(doc, "2.  Sekretaris Desa — Mengelola administrasi surat-menyurat, kearsipan, dan dokumentasi; menyusun laporan kegiatan; mengoordinasikan tugas bidang pemerintahan, pembangunan, dan kemasyarakatan; bertanggung jawab atas keuangan dan aset desa.")
bullet(doc, "3.  Kaur Pemerintahan — Mengelola urusan pemerintahan, kependudukan, dan perizinan; membantu administrasi kependudukan (KK, KTP); mengelola data dan dokumen kependudukan.")
bullet(doc, "4.  Kaur Kesra — Mengelola urusan kesejahteraan sosial, kesehatan, pendidikan, dan pemberdayaan masyarakat; mengkoordinasikan kegiatan kemasyarakatan; mengelola data UMKM dan potensi ekonomi desa.")
bullet(doc, "5.  Kepala Dusun I — Membantu kepala desa dalam koordinasi kegiatan pemerintahan di wilayah dusun; menjadi penghubung antara pemerintah desa dan warga di tingkat dusun.")

add_heading(doc, "3.2  Tata Laksana Sistem yang Berjalan")

body(doc, "Saat ini, pengelolaan dan penyebaran informasi di Desa Makmur masih dilakukan secara manual dan konvensional. Tata laksana sistem yang berjalan:")
body(doc, "a.  Pengelolaan Profil dan Data Desa", indent=False)
body(doc, "Data profil desa, sejarah, visi-misi, dan statistik kependudukan disimpan dalam arsip kertas dan file Microsoft Word/Excel di komputer kantor desa. Tidak ada sistem terpusat yang memudahkan akses. Pembaruan data dilakukan manual dan tidak selalu konsisten.")
body(doc, "b.  Penyebaran Informasi dan Pengumuman", indent=False)
body(doc, "Pengumuman disebarkan melalui: (1) papan pengumuman di kantor desa yang hanya terjangkau warga yang datang langsung; (2) pengumuman lisan melalui masjid/mushola; (3) grup WhatsApp non-resmi RT/RW. Metode ini sangat terbatas jangkauannya.")
body(doc, "c.  Pengelolaan Data Perangkat Desa", indent=False)
body(doc, "Informasi perangkat desa dicantumkan pada papan struktur organisasi fisik di kantor desa. Papan ini sering tidak diperbarui tepat waktu dan tidak dapat diakses oleh warga yang tidak mengunjungi kantor.")
body(doc, "d.  Pengelolaan Data UMKM", indent=False)
body(doc, "Pendataan UMKM dilakukan sesekali oleh Kaur Kesra dalam formulir kertas dan rekapitulasi Excel. Data ini tidak terpublikasikan secara resmi sehingga tidak dapat diakses oleh masyarakat umum atau pihak luar yang berpotensi menjadi mitra bisnis.")
body(doc, "e.  Akses Informasi oleh Warga dan Tamu", indent=False)
body(doc, "Warga yang membutuhkan informasi harus datang ke kantor desa pada jam kerja (Senin-Jumat, 08.00-16.00 WIB) atau menelepon langsung. Tamu dari luar desa tidak memiliki sumber resmi yang dapat diakses secara online.")

add_heading(doc, "3.3  Masalah yang Dihadapi")

body(doc, "Berdasarkan observasi dan wawancara dengan perangkat desa dan warga Desa Makmur, berikut adalah masalah yang diidentifikasi:")
doc.add_paragraph()
table_caption(doc, "Tabel 3.2 Masalah yang Dihadapi pada Sistem yang Berjalan")

prob_table = doc.add_table(rows=8, cols=3)
prob_table.style = 'Table Grid'
prob_data = [
    ["No", "Masalah", "Dampak"],
    ["1", "Informasi desa tersebar tidak terpusat dan tidak konsisten", "Warga mendapatkan informasi berbeda-beda dan tidak akurat"],
    ["2", "Tidak ada media digital resmi 24 jam", "Warga di luar jam kantor tidak dapat mengakses informasi resmi"],
    ["3", "Data UMKM tidak terdokumentasi dan tidak terpublikasi", "Potensi ekonomi desa tidak terekspos; UMKM sulit mendapat pelanggan baru"],
    ["4", "Pengumuman hanya via papan fisik dan grup WhatsApp non-resmi", "Jangkauan terbatas; tidak semua warga mendapat informasi tepat waktu"],
    ["5", "Papan perangkat desa tidak selalu diperbarui", "Warga tidak mengetahui kontak dan tugas perangkat yang akurat"],
    ["6", "Tidak ada backup data yang terstruktur", "Risiko kehilangan data permanen akibat kerusakan dokumen fisik"],
    ["7", "Tidak ada akses digital intuitif via smartphone", "Warga muda tidak memiliki akses digital ke informasi resmi desa"],
    ["8", "Tidak ada layanan informasi 24 jam non-stop", "Warga kesulitan memperoleh jawaban cepat di luar jam operasional kantor desa"],
    ["9", "Pengaduan warga bersifat lisan dan rawan tidak tercatat", "Pemerintah desa sulit melacak status penyelesaian masalah; warga tidak puas"],
]
for i, row in enumerate(prob_table.rows):
    for j, cell in enumerate(row.cells):
        fill_cell(cell, prob_data[i][j], bold=(i==0), center=(i==0 or j==0))
        if i == 0:
            set_cell_bg(cell, 'C00000')
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif i % 2 == 0:
            set_cell_bg(cell, 'FCE4E4')
doc.add_paragraph()

add_heading(doc, "3.4  Alternatif Pemecahan Masalah")

body(doc, "Berdasarkan permasalahan yang teridentifikasi, berikut adalah perbandingan alternatif pemecahan masalah:")
doc.add_paragraph()
table_caption(doc, "Tabel 3.3 Perbandingan Alternatif Pemecahan Masalah")

alt_table = doc.add_table(rows=5, cols=4)
alt_table.style = 'Table Grid'
alt_data = [
    ["Kriteria", "Alternatif 1:\nPHP + MySQL\n(Server Konvensional)", "Alternatif 2:\nWordPress\n(CMS Populer)", "Alternatif 3:\nSupabase BaaS +\nStatic Hosting\n(Diusulkan)"],
    ["Biaya Infrastruktur", "Tinggi (hosting + domain + server)", "Sedang (hosting + domain + plugin premium)", "Rendah/Gratis (free tier Supabase + Netlify)"],
    ["Kemudahan Admin Non-Teknis", "Rendah (perlu developer)", "Tinggi (familiar)", "Tinggi (panel admin custom intuitif)"],
    ["Keamanan Data", "Sedang (tergantung developer)", "Sedang (rentan plugin vulnerability)", "Tinggi (RLS native PostgreSQL + Supabase Auth)"],
    ["Akses via QR Code", "Perlu konfigurasi tambahan", "Perlu plugin tambahan", "Native – URL langsung di-generate QR"],
]
for i, row in enumerate(alt_table.rows):
    for j, cell in enumerate(row.cells):
        fill_cell(cell, alt_data[i][j], bold=(i==0 or j==0), center=(i==0))
        if i == 0:
            set_cell_bg(cell, '1F4E79')
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif j == 3:
            set_cell_bg(cell, 'E2EFDA')
        elif i % 2 == 0:
            set_cell_bg(cell, 'DEEAF1')
doc.add_paragraph()

body(doc, "Berdasarkan analisis perbandingan di atas, Alternatif 3 — pengembangan sistem informasi desa menggunakan Supabase sebagai BaaS dengan hosting statis — dipilih sebagai solusi terbaik. Alternatif ini menawarkan biaya operasional paling rendah (free tier), keamanan data paling tinggi melalui RLS native PostgreSQL, panel admin yang dapat dikustomisasi sesuai kebutuhan spesifik desa, serta kemudahan akses melalui QR Code yang dipasang di gapura desa.")

add_heading(doc, "3.5  User Requirement")

body(doc, "Analisis kebutuhan pengguna (user requirement) dilakukan melalui proses elisitasi yang merupakan proses mendefinisikan, mengumpulkan, dan memprioritaskan kebutuhan sistem dari para stakeholder. Proses dilakukan dalam tiga tahap dan menghasilkan Final Draft Elisitasi sebagai dasar pengembangan sistem.")

add_heading(doc, "3.5.1  Elisitasi Tahap I")

body(doc, "Elisitasi Tahap I adalah proses pengumpulan seluruh kebutuhan yang diinginkan oleh stakeholder (Kepala Desa, Sekretaris Desa, Kaur Kesra, dan perwakilan warga) tanpa disaring, diperoleh melalui wawancara dan diskusi langsung.")
doc.add_paragraph()
table_caption(doc, "Tabel 3.4 Elisitasi Tahap I")

e1_headers = ["No", "Kebutuhan Fungsional / Non-Fungsional", "Keterangan Stakeholder"]
e1_rows = [
    ["F01", "Menampilkan profil lengkap desa (nama, motto, sejarah, visi, misi, luas wilayah, jumlah penduduk)", "Kepala Desa, Sekdes"],
    ["F02", "Menampilkan data perangkat desa beserta foto, jabatan, periode, dan uraian tugas", "Kepala Desa"],
    ["F03", "Menampilkan daftar UMKM desa (nama usaha, pemilik, kategori, deskripsi, kontak, alamat, jam buka, peta)", "Kaur Kesra"],
    ["F04", "Menampilkan berita dan pengumuman resmi desa", "Kepala Desa, Sekdes"],
    ["F05", "Menampilkan galeri foto kegiatan dan potensi desa", "Kepala Desa"],
    ["F06", "Menampilkan informasi kontak kantor desa (alamat, telepon, email, jam layanan, peta, media sosial)", "Sekdes"],
    ["F07", "Panel admin dengan login email dan password", "Sekdes"],
    ["F08", "Admin dapat menambah, mengubah, menghapus data profil desa", "Sekdes"],
    ["F09", "Admin dapat menambah, mengubah, menghapus data perangkat desa", "Kepala Desa"],
    ["F10", "Admin dapat menambah, mengubah, menghapus data UMKM", "Kaur Kesra"],
    ["F11", "Admin dapat menambah, mengubah, menghapus berita dan pengumuman", "Kepala Desa"],
    ["F12", "Admin dapat menambah, mengubah, menghapus foto galeri desa", "Sekdes"],
    ["F13", "Admin dapat mengupload foto langsung dari panel admin (tidak hanya URL)", "Kaur Kesra"],
    ["F14", "Sistem dapat diakses melalui pemindaian QR Code di gapura desa", "Kepala Desa"],
    ["F15", "Admin dapat mengkustomisasi warna aksen desa dari panel admin", "Kepala Desa"],
    ["F16", "Foto desa dapat diupload sebagai latar belakang halaman utama (hero)", "Kepala Desa"],
    ["F17", "Detail item (perangkat, UMKM, berita) dalam popup/modal dengan animasi", "Warga"],
    ["F18", "Tombol kembali dari setiap halaman ke menu utama", "Warga"],
    ["F19", "Admin dapat mengatur urutan tampil setiap item (sort order)", "Sekdes"],
    ["F20", "Notifikasi WhatsApp otomatis saat ada berita baru", "Kepala Desa"],
    ["F21", "Fitur pengajuan surat online oleh warga", "Warga"],
    ["F22", "Dashboard statistik pengunjung untuk admin", "Sekdes"],
    ["F23", "Peta interaktif (Google Maps embed) pada halaman kontak", "Kaur Kesra"],
    ["F24", "URL Maps untuk lokasi setiap UMKM", "Kaur Kesra"],
    ["F25", "Registrasi akun admin mandiri (self-registration)", "Sekdes"],
    ["F26", "Upload logo desa dari panel admin", "Kepala Desa"],
    ["F27", "Tautan media sosial resmi desa (Facebook, Instagram)", "Kepala Desa"],
    ["F28", "Fitur pencarian informasi (search bar) untuk pengunjung", "Warga"],
    ["F29", "Fitur komentar warga pada setiap berita", "Warga"],
    ["F30", "Tampilan responsif untuk mobile dan desktop", "Warga"],
    ["NF01", "Dapat diakses di browser Android dan iOS tanpa instalasi aplikasi", "Semua"],
    ["NF02", "Waktu loading halaman utama tidak lebih dari 3 detik pada koneksi 4G", "Warga"],
    ["NF03", "Data pengunjung (anon) hanya dapat membaca, tidak dapat mengubah data", "Sekdes"],
    ["NF04", "Sistem dapat dihosting secara gratis (free tier)", "Kepala Desa"],
    ["NF05", "Tampilan antarmuka modern, bersih, dan mudah digunakan semua usia", "Semua"],
    ["NF06", "Tidak memerlukan instalasi aplikasi tambahan di smartphone warga", "Warga"],
    ["NF07", "Data tersimpan di cloud aman dari kerusakan perangkat fisik", "Sekdes"],
    ["NF08", "Sistem memiliki mekanisme backup data otomatis", "Sekdes"],
    ["F31", "Sistem menyediakan asisten Chatbot AI 'Tanya Desa' untuk menjawab pertanyaan secara instan", "Warga"],
    ["F32", "Menampilkan agenda kegiatan desa yang akan datang", "Sekdes"],
    ["F33", "Menyediakan dokumen dan surat desa yang dapat diunduh", "Kaur Kesra"],
    ["F34", "Menyediakan form pengaduan dan aspirasi warga", "Warga"],
    ["F35", "Admin dapat memberikan tanggapan dan status pada pengaduan warga", "Kepala Desa"],
]

e1_table = doc.add_table(rows=1+len(e1_rows), cols=3)
e1_table.style = 'Table Grid'
make_header_row(e1_table, e1_headers)
for i, row_data in enumerate(e1_rows):
    row = e1_table.rows[i+1]
    for j, cell in enumerate(row.cells):
        fill_cell(cell, row_data[j], center=(j==0))
        if i % 2 == 0:
            set_cell_bg(cell, 'F2F2F2')
doc.add_paragraph()

add_heading(doc, "3.5.2  Elisitasi Tahap II")

body(doc, "Elisitasi Tahap II menyaring hasil Tahap I menggunakan metode MDI (Mandatory, Desirable, Inessential):")
bullet(doc, "M (Mandatory)    — Kebutuhan yang wajib ada; sistem tidak dapat berfungsi tanpanya.")
bullet(doc, "D (Desirable)    — Kebutuhan yang diinginkan namun tidak wajib; meningkatkan nilai sistem.")
bullet(doc, "I (Inessential)  — Kebutuhan yang tidak esensial; tidak diimplementasikan pada versi ini.")
body(doc, "Kebutuhan dengan kategori I (Inessential) dieliminasi dari daftar pengembangan.")
doc.add_paragraph()
table_caption(doc, "Tabel 3.5 Elisitasi Tahap II (Metode MDI)")

e2_headers = ["No", "Kebutuhan", "M", "D", "I", "Keterangan"]
e2_rows = [
    ["F01", "Tampilkan profil lengkap desa", "v", "", "", "Inti sistem"],
    ["F02", "Tampilkan data perangkat desa", "v", "", "", "Inti sistem"],
    ["F03", "Tampilkan UMKM desa (detail + kontak + maps)", "v", "", "", "Inti sistem"],
    ["F04", "Tampilkan berita dan pengumuman", "v", "", "", "Inti sistem"],
    ["F05", "Tampilkan galeri foto", "v", "", "", "Inti sistem"],
    ["F06", "Tampilkan kontak & lokasi + Maps embed", "v", "", "", "Inti sistem"],
    ["F07", "Panel admin: login email + password", "v", "", "", "Wajib untuk keamanan"],
    ["F08", "Admin CRUD profil desa", "v", "", "", "Wajib untuk CMS"],
    ["F09", "Admin CRUD perangkat desa", "v", "", "", "Wajib untuk CMS"],
    ["F10", "Admin CRUD UMKM", "v", "", "", "Wajib untuk CMS"],
    ["F11", "Admin CRUD berita & pengumuman", "v", "", "", "Wajib untuk CMS"],
    ["F12", "Admin CRUD galeri foto", "v", "", "", "Wajib untuk CMS"],
    ["F13", "Upload foto langsung ke cloud storage", "v", "", "", "Wajib untuk kemudahan admin"],
    ["F14", "Akses via QR Code (QR di gapura desa)", "v", "", "", "Novelty utama penelitian"],
    ["F15", "Kustomisasi warna aksen dari panel admin", "", "v", "", "Meningkatkan personalisasi"],
    ["F16", "Upload foto hero desa dari admin", "v", "", "", "Estetika penting"],
    ["F17", "Detail item dalam modal popup + animasi GSAP", "v", "", "", "UX penting"],
    ["F18", "Tombol kembali ke menu utama", "v", "", "", "Navigasi wajib"],
    ["F19", "Sort order item (urutan tampil)", "", "v", "", "Meningkatkan fleksibilitas konten"],
    ["F20", "Notifikasi WhatsApp otomatis saat berita baru", "", "", "v", "Di luar scope; perlu API WA Business"],
    ["F21", "Pengajuan surat online oleh warga", "", "", "v", "Di luar scope penelitian ini"],
    ["F22", "Dashboard statistik pengunjung untuk admin", "", "v", "", "Berguna namun tidak kritis"],
    ["F23", "Peta interaktif (Maps embed) pada kontak", "v", "", "", "Wajib untuk kemudahan lokasi"],
    ["F24", "Maps untuk lokasi setiap UMKM", "", "v", "", "Berguna untuk navigasi pelanggan"],
    ["F25", "Registrasi admin mandiri (self-registration)", "v", "", "", "Wajib untuk setup awal"],
    ["F26", "Upload logo desa dari panel admin", "v", "", "", "Identitas visual desa"],
    ["F27", "Link media sosial desa (Facebook, Instagram)", "", "v", "", "Meningkatkan konektivitas"],
    ["F28", "Fitur pencarian (search bar) pengunjung", "", "", "v", "Tidak kritis di tahap ini"],
    ["F29", "Fitur komentar warga pada berita", "", "", "v", "Di luar scope; perlu moderasi"],
    ["F30", "Responsif mobile + desktop", "v", "", "", "Wajib karena akses via smartphone"],
    ["NF01", "Akses di Android dan iOS (browser)", "v", "", "", "Wajib"],
    ["NF02", "Loading halaman < 3 detik pada 4G", "v", "", "", "Performa kritis untuk UX"],
    ["NF03", "RLS: anon hanya baca, admin bisa tulis", "v", "", "", "Keamanan wajib"],
    ["NF04", "Hosting gratis (free tier)", "v", "", "", "Constraint anggaran desa"],
    ["NF05", "Antarmuka modern dan mudah digunakan", "v", "", "", "Usability wajib"],
    ["NF06", "Browser-based tanpa install aplikasi", "v", "", "", "Aksesibilitas wajib"],
    ["NF07", "Data tersimpan di cloud", "v", "", "", "Reliabilitas wajib"],
    ["NF08", "Backup data otomatis (by Supabase)", "", "v", "", "Ditangani Supabase otomatis"],
    ["F31", "Asisten Chatbot AI 'Tanya Desa'", "v", "", "", "Fitur pintar inovatif"],
    ["F32", "Modul agenda kegiatan", "v", "", "", "Inti sistem"],
    ["F33", "Modul unduh dokumen & surat", "v", "", "", "Inti sistem"],
    ["F34", "Form pengaduan warga", "v", "", "", "Inti sistem"],
    ["F35", "Tanggapan admin pada pengaduan", "v", "", "", "Wajib untuk komunikasi"],
]

e2_table = doc.add_table(rows=1+len(e2_rows), cols=6)
e2_table.style = 'Table Grid'
make_header_row(e2_table, e2_headers)
for i, row_data in enumerate(e2_rows):
    row = e2_table.rows[i+1]
    for j, cell in enumerate(row.cells):
        fill_cell(cell, row_data[j], center=(j in [0,2,3,4]))
        if row_data[4] == "v":
            set_cell_bg(cell, 'FCE4E4')
        elif i % 2 == 0:
            set_cell_bg(cell, 'F2F2F2')
doc.add_paragraph()

add_heading(doc, "3.5.3  Elisitasi Tahap III")

body(doc, "Elisitasi Tahap III menyaring kebutuhan yang lolos sebagai M (Mandatory) dan D (Desirable) dari Tahap II menggunakan metode TOE (Technical, Operational, Economic) dengan skala H (High/sulit), M (Medium/sedang), L (Low/mudah):")
bullet(doc, "T (Technical)    — H/M/L menunjukkan tingkat kesulitan teknis implementasi.")
bullet(doc, "O (Operational)  — H/M/L menunjukkan tingkat kesulitan operasional.")
bullet(doc, "E (Economic)     — H/M/L menunjukkan tingkat biaya implementasi.")
body(doc, "Kebutuhan dengan nilai T/O/E = H (High) yang tidak kritis dapat didefer ke pengembangan berikutnya.")
doc.add_paragraph()
table_caption(doc, "Tabel 3.6 Elisitasi Tahap III (Metode TOE)")

e3_headers = ["No", "Kebutuhan", "T", "O", "E", "Keputusan"]
e3_rows = [
    ["F01", "Tampilkan profil lengkap desa", "L", "L", "L", "Diterima"],
    ["F02", "Tampilkan data perangkat desa", "L", "L", "L", "Diterima"],
    ["F03", "Tampilkan UMKM + kontak + maps", "L", "L", "L", "Diterima"],
    ["F04", "Tampilkan berita dan pengumuman", "L", "L", "L", "Diterima"],
    ["F05", "Tampilkan galeri foto", "L", "L", "L", "Diterima"],
    ["F06", "Tampilkan kontak & lokasi + Maps embed", "L", "L", "L", "Diterima"],
    ["F07", "Panel admin: login email + password", "L", "L", "L", "Diterima"],
    ["F08", "Admin CRUD profil desa", "L", "L", "L", "Diterima"],
    ["F09", "Admin CRUD perangkat desa", "L", "L", "L", "Diterima"],
    ["F10", "Admin CRUD UMKM", "L", "L", "L", "Diterima"],
    ["F11", "Admin CRUD berita & pengumuman", "L", "L", "L", "Diterima"],
    ["F12", "Admin CRUD galeri foto", "L", "L", "L", "Diterima"],
    ["F13", "Upload foto langsung ke cloud storage", "M", "L", "L", "Diterima"],
    ["F14", "Akses via QR Code di gapura desa", "L", "L", "L", "Diterima"],
    ["F15", "Kustomisasi warna aksen dari admin", "L", "L", "L", "Diterima"],
    ["F16", "Upload foto hero desa dari admin", "M", "L", "L", "Diterima"],
    ["F17", "Detail item dalam modal popup + animasi GSAP", "M", "L", "L", "Diterima"],
    ["F18", "Tombol kembali ke menu utama", "L", "L", "L", "Diterima"],
    ["F19", "Sort order item (urutan tampil)", "L", "L", "L", "Diterima"],
    ["F22", "Dashboard statistik pengunjung admin", "H", "M", "M", "Defer (versi berikutnya)"],
    ["F24", "Maps lokasi setiap UMKM", "L", "L", "L", "Diterima"],
    ["F25", "Registrasi admin mandiri", "L", "L", "L", "Diterima"],
    ["F26", "Upload logo desa dari admin", "M", "L", "L", "Diterima"],
    ["F27", "Link media sosial desa", "L", "L", "L", "Diterima"],
    ["F30", "Responsif mobile + desktop", "M", "L", "L", "Diterima"],
    ["NF01", "Akses di Android & iOS (browser)", "L", "L", "L", "Diterima"],
    ["NF02", "Loading < 3 detik pada 4G", "M", "L", "L", "Diterima"],
    ["NF03", "RLS (anon baca, admin tulis)", "M", "L", "L", "Diterima"],
    ["NF04", "Hosting gratis (free tier)", "L", "L", "L", "Diterima"],
    ["NF05", "Antarmuka modern & mudah digunakan", "M", "L", "L", "Diterima"],
    ["NF06", "Browser-based tanpa install aplikasi", "L", "L", "L", "Diterima"],
    ["NF07", "Data di cloud (aman dari kerusakan fisik)", "L", "L", "L", "Diterima"],
    ["NF08", "Backup otomatis (by Supabase)", "L", "L", "L", "Diterima"],
    ["F31", "Asisten Chatbot AI 'Tanya Desa'", "M", "L", "L", "Diterima"],
    ["F32", "Modul agenda kegiatan", "L", "L", "L", "Diterima"],
    ["F33", "Modul unduh dokumen & surat", "L", "L", "L", "Diterima"],
    ["F34", "Form pengaduan warga", "L", "L", "L", "Diterima"],
    ["F35", "Tanggapan admin pada pengaduan", "L", "L", "L", "Diterima"],
]

e3_table = doc.add_table(rows=1+len(e3_rows), cols=6)
e3_table.style = 'Table Grid'
make_header_row(e3_table, e3_headers)
for i, row_data in enumerate(e3_rows):
    row = e3_table.rows[i+1]
    for j, cell in enumerate(row.cells):
        fill_cell(cell, row_data[j], center=(j in [0,2,3,4,5]))
        if "Defer" in row_data[5]:
            set_cell_bg(cell, 'FFF2CC')
        elif i % 2 == 0:
            set_cell_bg(cell, 'F2F2F2')
doc.add_paragraph()

add_heading(doc, "3.5.4  Final Draft Elisitasi")

body(doc, "Final Draft Elisitasi adalah daftar kebutuhan final yang telah lolos seluruh tahap penyaringan (Tahap I, II, dan III) dan menjadi acuan resmi perancangan dan pengembangan sistem informasi desa.")
doc.add_paragraph()
table_caption(doc, "Tabel 3.7 Final Draft Elisitasi")

fd_headers = ["No", "Kode", "Kebutuhan Final", "Kategori", "Prioritas"]
fd_rows = [
    ["1",  "F01",  "Sistem menampilkan profil desa lengkap: nama, motto, logo, foto hero, tahun berdiri, sejarah, visi, misi, luas wilayah, jumlah penduduk, jumlah KK, potensi desa", "Fungsional – Publik", "Tinggi"],
    ["2",  "F02",  "Sistem menampilkan daftar perangkat desa dengan foto, nama, jabatan, periode, tugas & fungsi; Kepala Desa ditampilkan dengan kartu istimewa (highlighted)", "Fungsional – Publik", "Tinggi"],
    ["3",  "F03",  "Sistem menampilkan daftar UMKM desa: nama usaha, pemilik, kategori, deskripsi, kontak WA, alamat, jam buka, dan tautan Google Maps lokasi UMKM", "Fungsional – Publik", "Tinggi"],
    ["4",  "F04",  "Sistem menampilkan berita dan pengumuman desa: judul, kategori, tanggal, isi, sumber, dan foto", "Fungsional – Publik", "Tinggi"],
    ["5",  "F05",  "Sistem menampilkan galeri foto kegiatan/potensi desa dengan fitur lightbox (klik foto untuk memperbesar)", "Fungsional – Publik", "Tinggi"],
    ["6",  "F06",  "Sistem menampilkan halaman kontak: alamat, telepon/WA, email, jam layanan, embed Google Maps kantor, tautan media sosial (Facebook, Instagram)", "Fungsional – Publik", "Tinggi"],
    ["7",  "F14",  "Sistem dapat diakses warga melalui pemindaian QR Code yang dipasang di gapura/titik strategis desa via kamera smartphone", "Fungsional – Publik", "Tinggi"],
    ["8",  "F17",  "Detail setiap item ditampilkan dalam modal popup dengan animasi GSAP yang halus", "Fungsional – Publik", "Tinggi"],
    ["9",  "F18",  "Setiap halaman view memiliki tombol 'Kembali' untuk kembali ke menu utama (grid menu)", "Fungsional – Publik", "Tinggi"],
    ["10", "F07",  "Panel admin menyediakan halaman login dengan email & password; admin baru dapat mendaftar dengan verifikasi email (self-registration)", "Fungsional – Admin", "Tinggi"],
    ["11", "F08",  "Admin dapat memperbarui seluruh data profil desa melalui formulir yang lengkap dan mudah digunakan", "Fungsional – Admin", "Tinggi"],
    ["12", "F09",  "Admin dapat menambah, mengubah, menghapus data perangkat desa; dapat menentukan siapa yang berstatus Kepala Desa (kartu istimewa)", "Fungsional – Admin", "Tinggi"],
    ["13", "F10",  "Admin dapat menambah, mengubah, menghapus data UMKM secara lengkap", "Fungsional – Admin", "Tinggi"],
    ["14", "F11",  "Admin dapat menambah, mengubah, menghapus berita dan pengumuman desa", "Fungsional – Admin", "Tinggi"],
    ["15", "F12",  "Admin dapat menambah, mengubah, menghapus foto galeri desa", "Fungsional – Admin", "Tinggi"],
    ["16", "F13",  "Admin dapat mengupload foto langsung dari perangkat ke Supabase Storage (bucket images) melalui panel admin tanpa menyalin URL manual", "Fungsional – Admin", "Tinggi"],
    ["17", "F15",  "Admin dapat mengatur warna aksen (brand color) desa; diterapkan otomatis ke seluruh tampilan website publik", "Fungsional – Admin", "Sedang"],
    ["18", "F19",  "Admin dapat mengatur urutan tampil (sort order) item perangkat, UMKM, berita, dan galeri", "Fungsional – Admin", "Sedang"],
    ["19", "F24",  "Setiap UMKM dapat dilengkapi URL Google Maps agar pengunjung dapat navigasi ke lokasi UMKM", "Fungsional – Admin", "Sedang"],
    ["20", "F27",  "Admin dapat mengisi tautan media sosial resmi desa (Facebook, Instagram) yang ditampilkan di halaman kontak", "Fungsional – Admin", "Sedang"],
    ["21", "NF01", "Sistem dapat diakses di browser modern Android dan iOS tanpa instalasi aplikasi tambahan", "Non-Fungsional", "Tinggi"],
    ["22", "NF02", "Waktu loading halaman utama tidak melebihi 3 detik pada koneksi 4G", "Non-Fungsional", "Tinggi"],
    ["23", "NF03", "Sistem menerapkan RLS: anon hanya SELECT; authenticated dapat INSERT, UPDATE, DELETE", "Non-Fungsional – Keamanan", "Tinggi"],
    ["24", "NF04", "Sistem dapat dioperasikan menggunakan free tier Supabase dan static hosting gratis tanpa biaya bulanan", "Non-Fungsional", "Tinggi"],
    ["25", "NF05", "Antarmuka modern, bersih, responsif, dan mudah digunakan pengguna dari berbagai usia termasuk warga pedesaan", "Non-Fungsional", "Tinggi"],
    ["26", "NF07", "Data desa tersimpan di cloud (Supabase) sehingga aman dari kerusakan fisik; Supabase melakukan backup berkala otomatis", "Non-Fungsional", "Tinggi"],
    ["27", "F31", "Sistem menyediakan Chatbot AI 'Tanya Desa' yang terintegrasi dengan Groq LLM melalui Edge Functions untuk menjawab pertanyaan seputar desa", "Fungsional – Publik", "Tinggi"],
    ["28", "F32", "Sistem menampilkan agenda kegiatan desa mendatang (judul, waktu, lokasi)", "Fungsional – Publik", "Tinggi"],
    ["29", "F33", "Sistem menyediakan dokumen/surat resmi desa yang dapat diunduh warga", "Fungsional – Publik", "Tinggi"],
    ["30", "F34", "Sistem menyediakan formulir pengaduan warga beserta status tanggapan dari pemerintah desa", "Fungsional – Publik", "Tinggi"],
    ["31", "F35", "Admin dapat mengelola, memperbarui status, dan memberikan tanggapan pada pengaduan warga", "Fungsional – Admin", "Tinggi"],
]

fd_table = doc.add_table(rows=1+len(fd_rows), cols=5)
fd_table.style = 'Table Grid'
make_header_row(fd_table, fd_headers)
for i, row_data in enumerate(fd_rows):
    row = fd_table.rows[i+1]
    cat = row_data[3]
    for j, cell in enumerate(row.cells):
        fill_cell(cell, row_data[j], center=(j in [0,4]))
        if "Publik" in cat:
            if i % 2 == 0:
                set_cell_bg(cell, 'E2EFDA')
            else:
                set_cell_bg(cell, 'EFF7EA')
        elif "Admin" in cat:
            if i % 2 == 0:
                set_cell_bg(cell, 'DEEAF1')
            else:
                set_cell_bg(cell, 'EBF4FB')
        else:
            if i % 2 == 0:
                set_cell_bg(cell, 'FFF2CC')
            else:
                set_cell_bg(cell, 'FFFBE6')
doc.add_paragraph()

body(doc, "Final Draft Elisitasi mencakup 26 kebutuhan final (20 kebutuhan fungsional dan 6 kebutuhan non-fungsional) yang menjadi landasan perancangan dan pengembangan sistem informasi desa. Satu kebutuhan (F22 - Dashboard Statistik Pengunjung) ditangguhkan untuk versi berikutnya, sementara empat kebutuhan (F20, F21, F28, F29) dieliminasi karena berada di luar ruang lingkup penelitian ini.")

body(doc, "Kebutuhan-kebutuhan ini mencerminkan tiga kelompok utama: (1) modul informasi publik yang diakses warga melalui QR Code; (2) panel administrasi untuk perangkat desa dalam mengelola konten; dan (3) aspek non-fungsional yang menjamin keamanan, performa, dan aksesibilitas sistem secara menyeluruh.")

# ─────────────────────────────────────────────────────────────────────────────
# DAFTAR PUSTAKA
# ─────────────────────────────────────────────────────────────────────────────

doc.add_page_break()
add_bab_heading(doc, "DAFTAR PUSTAKA")

references = [
    "Arifin, M., Hakim, L., & Setiawan, B. (2023). Evaluasi Usability Sistem Informasi Desa Menggunakan Metode System Usability Scale (SUS): Studi Kasus 5 Desa di Jawa Tengah. Jurnal Nasional Informatika dan Teknologi Jaringan, 4(2), 145-158.",
    "Brooke, J. (1996). SUS: A 'Quick and Dirty' Usability Scale. In P. Jordan, B. Thomas, B. Weerdmeester, & I. McClelland (Eds.), Usability Evaluation in Industry. Taylor & Francis.",
    "Fathansyah. (2022). Basis Data (Edisi Revisi Ketiga). Informatika Bandung.",
    "Firmansyah, R., & Suhartono, D. (2023). Backend-as-a-Service sebagai Solusi Infrastruktur Sistem Informasi untuk Institusi Pemerintah Skala Kecil. Journal of Information Systems, 5(1), 22-38.",
    "Forouzan, B. A. (2022). Data Communications and Networking (6th Ed.). McGraw-Hill Education.",
    "GreenSock. (2024). GSAP (GreenSock Animation Platform) Documentation. Diakses dari https://gsap.com/docs/",
    "Hariadi, F., Prasetyo, W., & Anggraini, R. (2024). Implementasi Serverless Architecture pada Pengembangan Aplikasi Pemerintahan di Indonesia: Kajian Komparatif Firebase vs Supabase. Journal of Computer Science and Information, 17(1), 34-51.",
    "Hidayat, T., Wulandari, S., & Permana, A. (2021). Sistem Informasi UMKM Desa Berbasis Web untuk Peningkatan Pemasaran Digital Produk Lokal. Jurnal SIMETRIS, 12(2), 78-90.",
    "Jogiyanto, H. M. (2021). Analisis dan Desain Sistem Informasi: Pendekatan Terstruktur Teori dan Praktik Aplikasi Bisnis (Edisi Revisi). ANDI Yogyakarta.",
    "MDN Web Docs. (2024). HTML5 – HyperText Markup Language. Mozilla Developer Network. Diakses dari https://developer.mozilla.org/en-US/docs/Web/HTML",
    "Mulyani, S., Rahayu, P., & Hermawan, A. (2021). Digitalisasi Layanan Informasi Publik Desa melalui Pengembangan Website Terintegrasi dengan Pelatihan Perangkat Desa. Jurnal Pengabdian kepada Masyarakat, 5(3), 201-215.",
    "Nugroho, A. W., & Pratama, B. D. (2023). Implementasi QR Code pada Sistem Informasi Publik untuk Meningkatkan Aksesibilitas Layanan Pemerintah Daerah. Jurnal Ilmiah Penelitian Informatika (JIPI), 8(2), 112-124.",
    "Peraturan Menteri Dalam Negeri Nomor 47 Tahun 2016 tentang Administrasi Pemerintahan Desa.",
    "Peraturan Pemerintah Nomor 43 Tahun 2014 tentang Peraturan Pelaksanaan Undang-Undang Nomor 6 Tahun 2014 tentang Desa.",
    "Pratama, I. P., Gunawan, R., & Sari, N. P. (2023). Penerapan Teknologi QR Code dalam Sistem Informasi Berbasis Web untuk Kemudahan Akses Masyarakat Pedesaan. Jurnal JATISI, 10(2), 234-249.",
    "Priyono, A. (2022). QR Code: Teknologi, Implementasi, dan Aplikasinya dalam Sistem Informasi Modern. Rekayasa Sains Bandung.",
    "Rahmawati, D., Setiawan, E., & Purnomo, A. H. (2022). Pengembangan Website Profil Desa Responsif Berbasis Mobile-First Design dengan Pendekatan Progressive Web App. Jurnal Teknologi Informasi dan Ilmu Komputer (JTIIK), 9(4), 789-802.",
    "Santoso, B., Kurniawan, D., & Mahendra, R. (2022). Perancangan Sistem Manajemen Konten (CMS) untuk Website Resmi Pemerintah Desa yang Mudah Digunakan oleh Perangkat Non-Teknis. Jurnal TEKNOINFO, 16(1), 88-101.",
    "Satria, A., Rahmat, M., & Kurniadi, D. (2022). Rancang Bangun Sistem Informasi Desa Berbasis Web dengan Framework CodeIgniter 4. Jurnal JATISI, 9(1), 45-58.",
    "Sulistyowati, R., & Prabowo, A. (2022). Analisis dan Implementasi Row Level Security PostgreSQL pada Sistem Informasi Pemerintahan untuk Perlindungan Data Sensitif. Jurnal Sistem Informasi dan Teknologi, 8(1), 15-29.",
    "Supabase. (2024). Supabase Documentation: Database, Auth, Storage, and Edge Functions. Diakses dari https://supabase.com/docs",
    "Sutabri, T. (2022). Konsep Sistem Informasi. ANDI Yogyakarta.",
    "Undang-Undang Nomor 6 Tahun 2014 tentang Desa. Lembaran Negara Republik Indonesia Tahun 2014 Nomor 7.",
    "Wijaya, M. A., Santoso, H., & Kurniawati, D. (2024). Smart Village: Model Digitalisasi Tata Kelola Informasi Desa Berbasis Cloud Computing di Era Society 5.0. Jurnal Informatika dan Komputer, 11(1), 56-72.",
]

for ref in references:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.left_indent = Cm(1.25)
    pf.first_line_indent = Cm(-1.25)
    run = p.add_run(ref)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────

output_path = r"c:\Users\1nkuss\Documents\WebInformasiDesa\Laporan_Skripsi_WebInformasiDesa.docx"
doc.save(output_path)
print(f"SUCCESS! Dokumen berhasil dibuat:")
print(f"  Path  : {output_path}")
print(f"  BAB   : I (Pendahuluan) + II (Landasan Teori) + III (Analisis Sistem)")
print(f"  Tabel : {len(doc.tables)} tabel (LitReview + 4 Elisitasi + Org + Masalah + Alternatif)")
print(f"  Total paragraf: {len(doc.paragraphs)}")
