"""
Script untuk menambahkan Daftar Simbol ke laporan skripsi WebInformasiDesa.
Daftar Simbol disisipkan SETELAH Bab 2 (Landasan Teori) dan SEBELUM Bab 3 (Analisis Sistem).
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─────────────────────────────────────────────────────────────────────────────
# Data simbol yang akan ditampilkan
# ─────────────────────────────────────────────────────────────────────────────
SYMBOLS = [
    # ── Use Case Diagram ──────────────────────────────────────────────────────
    {
        "kategori": "Use Case Diagram",
        "simbol": "●  (Aktor)",
        "nama": "Aktor (Actor)",
        "keterangan": "Mewakili pengguna atau sistem eksternal yang berinteraksi dengan sistem yang dimodelkan.",
    },
    {
        "kategori": "",
        "simbol": "○  (Use Case)",
        "nama": "Use Case",
        "keterangan": "Mewakili fungsionalitas atau layanan yang disediakan oleh sistem kepada aktor.",
    },
    {
        "kategori": "",
        "simbol": "─────",
        "nama": "Asosiasi (Association)",
        "keterangan": "Menunjukkan hubungan komunikasi antara aktor dan use case.",
    },
    {
        "kategori": "",
        "simbol": "<<include>>",
        "nama": "Include",
        "keterangan": "Menunjukkan bahwa satu use case selalu menyertakan perilaku use case lain.",
    },
    {
        "kategori": "",
        "simbol": "<<extend>>",
        "nama": "Extend",
        "keterangan": "Menunjukkan bahwa satu use case dapat memperluas perilaku use case lain secara kondisional.",
    },
    {
        "kategori": "",
        "simbol": "□  (Sistem)",
        "nama": "Batas Sistem (System Boundary)",
        "keterangan": "Kotak persegi panjang yang mendefinisikan lingkup atau batas sistem yang dimodelkan.",
    },
    # ── Activity Diagram ──────────────────────────────────────────────────────
    {
        "kategori": "Activity Diagram",
        "simbol": "⬤  (Awal)",
        "nama": "Initial Node (Start)",
        "keterangan": "Titik awal alur aktivitas, digambarkan sebagai lingkaran hitam penuh.",
    },
    {
        "kategori": "",
        "simbol": "◉  (Akhir)",
        "nama": "Activity Final Node (End)",
        "keterangan": "Titik akhir alur aktivitas, digambarkan sebagai lingkaran hitam penuh dengan cincin.",
    },
    {
        "kategori": "",
        "simbol": "▭  (Aktivitas)",
        "nama": "Action / Activity",
        "keterangan": "Persegi panjang dengan sudut melengkung yang merepresentasikan satu langkah proses atau aktivitas.",
    },
    {
        "kategori": "",
        "simbol": "◇  (Keputusan)",
        "nama": "Decision Node",
        "keterangan": "Belah ketupat yang menunjukkan percabangan alur berdasarkan kondisi tertentu (ya/tidak).",
    },
    {
        "kategori": "",
        "simbol": "═══  (Fork/Join)",
        "nama": "Fork / Join",
        "keterangan": "Batang hitam tebal yang menunjukkan percabangan (fork) atau penggabungan (join) alur paralel.",
    },
    {
        "kategori": "",
        "simbol": "→  (Alur)",
        "nama": "Control Flow",
        "keterangan": "Anak panah yang menunjukkan arah alur kontrol dari satu aktivitas ke aktivitas berikutnya.",
    },
    {
        "kategori": "",
        "simbol": "║  (Swimlane)",
        "nama": "Swimlane",
        "keterangan": "Pembatas vertikal atau horizontal yang memisahkan aktivitas berdasarkan aktor atau sistem yang bertanggung jawab.",
    },
    # ── ERD (Entity Relationship Diagram) ────────────────────────────────────
    {
        "kategori": "Entity Relationship Diagram (ERD)",
        "simbol": "▭  (Entitas)",
        "nama": "Entitas (Entity)",
        "keterangan": "Persegi panjang yang merepresentasikan objek atau konsep nyata yang memiliki data tersimpan.",
    },
    {
        "kategori": "",
        "simbol": "◇  (Relasi)",
        "nama": "Relasi (Relationship)",
        "keterangan": "Belah ketupat yang menunjukkan hubungan antara dua entitas atau lebih.",
    },
    {
        "kategori": "",
        "simbol": "○  (Atribut)",
        "nama": "Atribut (Attribute)",
        "keterangan": "Elips yang merepresentasikan properti atau karakteristik dari sebuah entitas.",
    },
    {
        "kategori": "",
        "simbol": "○̲  (PK)",
        "nama": "Atribut Kunci (Primary Key)",
        "keterangan": "Elips dengan nama yang digarisbawahi, merepresentasikan atribut unik pengidentifikasi setiap record entitas.",
    },
    {
        "kategori": "",
        "simbol": "1, N, M",
        "nama": "Kardinalitas",
        "keterangan": "Notasi yang menunjukkan jumlah minimum dan maksimum keterlibatan entitas dalam relasi: 1 (satu), N (banyak), M (banyak).",
    },
    # ── Flowchart / Kamus Data ────────────────────────────────────────────────
    {
        "kategori": "Flowchart & Kamus Data",
        "simbol": "▱  (Proses)",
        "nama": "Proses",
        "keterangan": "Persegi panjang yang menunjukkan suatu proses atau operasi komputasi/pengolahan data.",
    },
    {
        "kategori": "",
        "simbol": "⬡  (Terminal)",
        "nama": "Terminal (Terminator)",
        "keterangan": "Bentuk oval/pill yang menandai titik awal (Start) atau akhir (End) dari sebuah alur proses.",
    },
    {
        "kategori": "",
        "simbol": "▱  (I/O)",
        "nama": "Input / Output",
        "keterangan": "Jajaran genjang yang merepresentasikan operasi masukan (input) atau keluaran (output) data.",
    },
    {
        "kategori": "",
        "simbol": "∗  (Iterasi)",
        "nama": "Iterasi (Pengulangan)",
        "keterangan": "Digunakan dalam kamus data untuk menandai elemen data yang dapat berulang.",
    },
    {
        "kategori": "",
        "simbol": "+  (Dan)",
        "nama": "Dan (AND)",
        "keterangan": "Digunakan dalam kamus data untuk menghubungkan beberapa elemen yang kesemuanya harus ada.",
    },
    {
        "kategori": "",
        "simbol": "|  (Pilihan)",
        "nama": "Pilihan (OR)",
        "keterangan": "Digunakan dalam kamus data untuk menunjukkan salah satu dari beberapa elemen yang harus dipilih.",
    },
    # ── Notasi Pengujian (SUS) ────────────────────────────────────────────────
    {
        "kategori": "Notasi Pengujian & Evaluasi (SUS)",
        "simbol": "X̄",
        "nama": "Rata-rata (Mean)",
        "keterangan": "Nilai rata-rata skor yang diperoleh dari keseluruhan responden dalam kuesioner SUS.",
    },
    {
        "kategori": "",
        "simbol": "Σ",
        "nama": "Sigma (Jumlah)",
        "keterangan": "Notasi matematis untuk penjumlahan seluruh nilai dalam suatu himpunan data.",
    },
    {
        "kategori": "",
        "simbol": "n",
        "nama": "Jumlah Responden",
        "keterangan": "Jumlah total responden atau sampel yang digunakan dalam pengujian.",
    },
    {
        "kategori": "",
        "simbol": "Si",
        "nama": "Skor Item ke-i",
        "keterangan": "Skor yang diberikan responden untuk pertanyaan ke-i pada kuesioner SUS.",
    },
    {
        "kategori": "",
        "simbol": "SSUS",
        "nama": "Skor SUS",
        "keterangan": "Nilai akhir System Usability Scale dalam skala 0–100 yang diperoleh dari konversi skor mentah kuesioner.",
    },
]

# ─────────────────────────────────────────────────────────────────────────────
# Fungsi bantu
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set background color of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_col_width(table, col_widths_cm):
    """Set column widths by modifying the XML directly."""
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(col_widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement("w:tcW")
                # 1 cm ≈ 567 twips
                tcW.set(qn("w:w"), str(int(col_widths_cm[i] * 567)))
                tcW.set(qn("w:type"), "dxa")
                tcPr.append(tcW)


def make_bold_run(para, text, size_pt=10):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size_pt)
    return run


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def build_symbol_doc(source_path: str, output_path: str):
    doc = Document(source_path)

    # ── 1. Temukan paragraf "BAB III" ──────────────────────────────────────
    bab3_idx = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().startswith("BAB III"):
            bab3_idx = i
            break

    if bab3_idx is None:
        print("⚠  BAB III tidak ditemukan, simbol akan ditambahkan di akhir dokumen.")
        bab3_idx = len(doc.paragraphs)

    print(f"✔  BAB III ditemukan di paragraf index {bab3_idx}")

    # ── 2. Sisipkan halaman Daftar Simbol ──────────────────────────────────
    # Kita akses parent XML element di antara paragraf bab2_last dan bab3_first
    body = doc.element.body

    # Ambil elemen XML paragraf ke-(bab3_idx - 1) sebagai anchor
    all_paras_xml = [c for c in body if c.tag.endswith("}p") or c.tag.endswith("}tbl")]
    
    # Temukan elemen XML yang sesuai dengan doc.paragraphs[bab3_idx]
    target_para_xml = doc.paragraphs[bab3_idx]._element

    def insert_before(anchor_xml, new_element):
        """Insert new_element immediately before anchor_xml in body."""
        anchor_xml.addprevious(new_element)

    # Helper: buat paragraf baru dengan style Normal
    def new_para(text="", bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, center=False):
        from docx.oxml.ns import qn as qname
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qname("w:val"), "Normal")
        pPr.append(pStyle)
        if center:
            jc = OxmlElement("w:jc")
            jc.set(qname("w:val"), "center")
            pPr.append(jc)
        p.append(pPr)
        if text:
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            if bold:
                b = OxmlElement("w:b")
                rPr.append(b)
            sz = OxmlElement("w:sz")
            sz.set(qname("w:val"), str(size * 2))
            rPr.append(sz)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.text = text
            if text.startswith(" ") or text.endswith(" "):
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            p.append(r)
        return p

    # ── 2a. Page break sebelum Daftar Simbol ──────────────────────────────
    pb_para = OxmlElement("w:p")
    pb_r = OxmlElement("w:r")
    pb_br = OxmlElement("w:br")
    pb_br.set(qn("w:type"), "page")
    pb_r.append(pb_br)
    pb_para.append(pb_r)
    insert_before(target_para_xml, pb_para)

    # ── 2b. Judul Halaman ──────────────────────────────────────────────────
    title_p = new_para("DAFTAR SIMBOL", bold=True, size=14, center=True)
    insert_before(target_para_xml, title_p)

    # Spasi setelah judul
    insert_before(target_para_xml, new_para(""))
    insert_before(target_para_xml, new_para(""))

    # ── 2c. Buat tabel Daftar Simbol ──────────────────────────────────────
    # Tabel langsung di dokumen, tapi kita harus manualkan XML karena
    # doc.add_table() tidak support insert before — kita pakai cara sah:
    # Buat tabel terpisah, lalu pindahkan XML-nya.

    tmp_doc = Document()
    tbl = tmp_doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ── Header row ────────────────────────────────────────────────────────
    hdr = tbl.rows[0].cells
    headers = ["No.", "Simbol", "Nama Simbol", "Keterangan"]
    col_widths = [1.0, 3.0, 4.5, 8.0]  # cm

    for i, (cell, hdr_text) in enumerate(zip(hdr, headers)):
        set_cell_bg(cell, "2E7D32")  # dark green
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(hdr_text)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

    # ── Data rows ─────────────────────────────────────────────────────────
    no = 1
    current_kategori = None

    for entry in SYMBOLS:
        row = tbl.add_row()
        cells = row.cells

        # Jika kategori berubah, tambahkan sub-heading row
        if entry["kategori"] and entry["kategori"] != current_kategori:
            current_kategori = entry["kategori"]

            # Merge seluruh kolom jadi 1 untuk sub-heading
            cat_row = tbl.add_row()
            cat_cells = cat_row.cells
            # Merge cells
            merged = cat_cells[0].merge(cat_cells[1]).merge(cat_cells[2]).merge(cat_cells[3])
            set_cell_bg(merged, "C8E6C9")  # light green
            p = merged.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"  {current_kategori}")
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)

            # Tambah data row setelahnya
            row = tbl.add_row()
            cells = row.cells

        # No.
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[0].paragraphs[0].add_run(str(no)).font.size = Pt(10)

        # Simbol
        cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cells[1].paragraphs[0].add_run(entry["simbol"])
        r.font.size = Pt(11)

        # Nama simbol
        cells[2].paragraphs[0].add_run(entry["nama"]).font.size = Pt(10)

        # Keterangan
        cells[3].paragraphs[0].add_run(entry["keterangan"]).font.size = Pt(10)

        # Zebra striping — baris genap diberi warna abu-abu sangat muda
        if no % 2 == 0:
            for c in cells:
                set_cell_bg(c, "F9FBE7")

        no += 1

    # Set column widths
    set_col_width(tbl, col_widths)

    # Pindahkan tbl XML ke dokumen utama, sebelum BAB III
    tbl_xml = tbl._tbl
    insert_before(target_para_xml, tbl_xml)

    # ── 2d. Spasi setelah tabel + keterangan singkat ───────────────────────
    insert_before(target_para_xml, new_para(""))
    note_p = new_para(
        "Catatan: Simbol-simbol di atas digunakan dalam diagram dan kamus data pada Bab III dan Bab IV laporan ini.",
        bold=False, size=9
    )
    insert_before(target_para_xml, note_p)
    insert_before(target_para_xml, new_para(""))

    # ── 3. Simpan dokumen ─────────────────────────────────────────────────
    doc.save(output_path)
    print(f"✅  Dokumen berhasil disimpan: {output_path}")
    print(f"    Total simbol ditambahkan: {sum(1 for e in SYMBOLS)}")


if __name__ == "__main__":
    build_symbol_doc(
        source_path="Laporan_Skripsi_WebInformasiDesa.docx",
        output_path="Laporan_Skripsi_WebInformasiDesa_v2.docx",
    )
